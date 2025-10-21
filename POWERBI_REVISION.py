"""
POWERBI_REVISION
"""

#!/usr/bin/env python
# coding: utf-8

# In[5]:


import os
import openai
import subprocess
from docx import Document
from dotenv import load_dotenv

# === Cargar configuración ===
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# === Ruta fija del archivo PBIX ===
PBIX_PATH = r"C:\Users\HP\Desktop\PROPUESTAS DE DIPLOMADO\DIPLOMADO EXCEL Y POWER BI APLICADO A LA GESTION DE VENTAS\CURSOS DE MI MODULO\TRABAJOS\SEMANA 2 ENTREBGABLE 1\BENJAMIN MARTINEZ MARTINEZ_75717_assignsubmission_file_\Trabajo_final_Benjamin_Martinez_Martinez.pbix"

# === Rutas para archivo temporal e informe ===
EXTRACTED_MODEL_PATH = os.path.join(os.path.dirname(PBIX_PATH), "modelo_extraido.txt")
OUTPUT_DOC_PATH = os.path.join(os.path.dirname(PBIX_PATH), "informe_auditoria_powerbi.docx")
TABULAR_EDITOR_CLI = r"C:\Program Files (x86)\Tabular Editor\TabularEditor.exe"  # Usa ruta completa si no está en el PATH

# === Función: Exportar modelo con Tabular Editor CLI ===
def exportar_modelo_tabular_editor(pbix_path: str, output_path: str):
    print("🔍 Exportando modelo con Tabular Editor...")
    comando = f'"{TABULAR_EDITOR_CLI}" "{pbix_path}" -S "ExportModel" -O "{output_path}"'
    resultado = subprocess.run(comando, shell=True)
    if resultado.returncode != 0:
        print("❌ Error al exportar el modelo.")
    else:
        print("✅ Modelo exportado correctamente.")

# === Función: Leer el modelo exportado ===
def leer_modelo(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

# === Crear prompts para OpenAI ===
def crear_prompt_modelo(texto_modelo: str) -> str:
    return f"""
Actúa como un auditor experto en Power BI. Evalúa lo siguiente:

1. Calidad del dimensionamiento (relaciones, cardinalidades, uso de dimensiones).
2. Calidad de las medidas DAX.
3. Buenas prácticas y errores comunes.
4. Recomendaciones de mejora profesional.

Modelo extraído:
{texto_modelo[:10000]}...
"""

def crear_prompt_visuales() -> str:
    return """
Evalúa visualizaciones típicas de Power BI según buenas prácticas:

- Claridad del mensaje visual.
- Uso adecuado de gráficos según los datos.
- Contraste, etiquetas, títulos, leyendas.
- Recomendaciones visuales generales para mejorar dashboards profesionales.

(No se incluyen imágenes en esta auditoría).
"""

# === Función: Consultar a OpenAI ===
def consultar_openai(prompt: str) -> str:
    print("💬 Consultando OpenAI...")
    respuesta = openai.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return respuesta.choices[0].message.content

# === Función: Generar el informe en Word ===
def generar_informe(titulo: str, hallazgos: dict, salida: str):
    print("📝 Generando informe Word...")
    doc = Document()
    doc.add_heading(titulo, 0)
    for seccion, contenido in hallazgos.items():
        doc.add_heading(seccion, level=1)
        doc.add_paragraph(contenido)
    doc.save(salida)
    print(f"✅ Informe guardado en: {salida}")

# === Función principal ===
def main():
    if not os.path.exists(PBIX_PATH):
        print(f"❌ No se encuentra el archivo PBIX: {PBIX_PATH}")
        return

    # Paso 1: Exportar modelo
    exportar_modelo_tabular_editor(PBIX_PATH, EXTRACTED_MODEL_PATH)

    # Paso 2: Leer modelo
    if not os.path.exists(EXTRACTED_MODEL_PATH):
        print("❌ No se pudo extraer el modelo.")
        return

    modelo_texto = leer_modelo(EXTRACTED_MODEL_PATH)

    # Paso 3: Preparar prompts
    prompt_modelo = crear_prompt_modelo(modelo_texto)
    prompt_visual = crear_prompt_visuales()

    # Paso 4: Consultar OpenAI
    resultado_modelo = consultar_openai(prompt_modelo)
    resultado_visual = consultar_openai(prompt_visual)

    # Paso 5: Generar informe
    hallazgos = {
        "Auditoría del Modelo de Datos": resultado_modelo,
        "Auditoría de Visualizaciones": resultado_visual
    }
    generar_informe("Informe de Auditoría Power BI", hallazgos, OUTPUT_DOC_PATH)

# === Ejecutar ===
if __name__ == "__main__":
    main()


# In[ ]:




